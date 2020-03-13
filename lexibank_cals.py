"""
Note: We run libreoffice to convert from doc to docx after download.
"""

import re
from collections import defaultdict
from pathlib import Path
from subprocess import check_call

from clldutils.misc import slug
from clldutils.path import Path
from csvw.dsv import UnicodeWriter, reader
from docx import Document
from pylexibank.dataset import Dataset as BaseDataset
from pylexibank.forms import FormSpec

SOURCE = "Mennecier2016"


class Dataset(BaseDataset):
    dir = Path(__file__).parent
    id = "cals"

    # split(" ~ ")
    form_spec = FormSpec(
        brackets={},
        separators="~",
        missing_data=(),
        strip_inside_brackets=False,
    )

    def cmd_download(self, args):
        fname = self.raw_dir / "Table_S2_Supplementary_Mennecier_et_al..doc"

        self.raw_dir.download_and_unpack(
            "https://ndownloader.figshare.com/articles/3443090/versions/1",
            fname,
            log=args.log,
        )

        check_call(
            "libreoffice --headless --convert-to docx %s --outdir %s"
            % (fname, self.raw_dir),
            shell=True,
        )

        doc = Document(
            self.raw_dir / "Table_S2_Supplementary_Mennecier_et_al..docx"
        )
        for i, table in enumerate(doc.tables):
            with UnicodeWriter(
                self.raw_dir.joinpath("%s.csv" % (i + 1,)).as_posix()
            ) as writer:
                for row in table.rows:
                    # This code fixes a wrong gloss in the raw source,
                    # where the `itʲerʊ` set is glossed as "to pull"
                    # instead of the correct "to push". See discussion
                    # at https://github.com/lexibank/cals/pull/7
                    row_data = map(text_and_color, row.cells)
                    if i == 11:
                        row_data = [
                            cell if cell != "to pull" else "to push"
                            for cell in row_data
                        ]
                    writer.writerow(row_data)

    def cmd_makecldf(self, args):
        gcode = {x["ID"]: x["Glottocode"] for x in self.languages}
        data = defaultdict(dict)
        args.writer.add_sources()

        for fname in sorted(self.raw_dir.glob("*.csv")):
            read(fname, data)

        ccode = args.writer.add_concepts(id_factory=lambda c: slug(c.label))

        # Add manual correction
        ccode.append("topush")
        args.writer.add_concept(
            ID="topush",
            Name="to push",
            Concepticon_ID="1452",
            Concepticon_Gloss="PUSH",
        )

        for doculect, wl in sorted(data.items()):
            sd = slug(doculect).capitalize()
            args.writer.add_language(
                ID=sd, Name=doculect, Glottocode=gcode[doculect.split("-")[0]]
            )

            for concept, (form, loan, cogset) in sorted(wl.items()):
                sc = slug(concept)
                if sc in ccode:
                    pass
                elif sc.startswith("to ") and sc[3:] in ccode:
                    sc = sc[3:]
                else:
                    sc = None

                for row in args.writer.add_lexemes(
                    Language_ID=sd, Parameter_ID=sc, Value=form, Source=SOURCE
                ):
                    if cogset:
                        args.writer.add_cognate(
                            lexeme=row,
                            Cognateset_ID="%s-%s" % (sc, slug(cogset)),
                        )
                        break


COLOR_PATTERN = re.compile('fill="(?P<color>[^"]+)"')


def text_and_color(cell):
    color = None
    for line in cell._tc.tcPr.xml.split("\n"):
        if "w:shd" in line:
            m = COLOR_PATTERN.search(line)
            if m:
                color = m.group("color")
                break
    if color == "auto":
        color = None
    if color:
        color = "#" + color + " "
    return "%s%s" % (color if color else "", cell.paragraphs[0].text)


def get_loan_and_form(c):
    if c.startswith("#"):
        return c.split(" ", 1)
    return None, c


def read(fname, data):
    concepts, loan = None, None

    for i, row in enumerate(reader(fname)):
        if i == 0:
            concepts = {j: c for j, c in enumerate(row[1:])}
        else:
            for j, c in enumerate(row[1:]):
                if j % 2 == 0:  # even number
                    loan, form = get_loan_and_form(c)
                else:
                    if form.strip():
                        data[row[0]][concepts[j]] = (form, loan, c)
    return data
